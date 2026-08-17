"""DOCX -> ordered raw elements.

`python-docx`'s `.paragraphs` and `.tables` properties lose interleaved
order (a table between two paragraphs), which this project must preserve.
This module walks the underlying OpenXML body directly to recover true
document order.
"""

from __future__ import annotations

from collections.abc import Iterator
from dataclasses import dataclass
from pathlib import Path
from typing import Union

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from loguru import logger


@dataclass(frozen=True)
class RawCell:
    """One table cell plus OpenXML merge metadata. `grid_span` covers
    horizontal merges; `v_merge` covers vertical merges ('restart' begins
    a merged region, 'continue' is a continuation cell).
    """

    text: str
    grid_span: int
    v_merge: str | None  # None | "restart" | "continue"


@dataclass(frozen=True)
class RawTable:
    rows: list[list[RawCell]]


@dataclass(frozen=True)
class RawParagraph:
    text: str
    style_name: str
    outline_level: int | None  # from paragraph properties, if a heading style sets one
    numbering_id: int | None  # non-None if the paragraph has explicit list/numbering


RawElement = Union[RawParagraph, RawTable]


def _iter_block_items(document: DocxDocument) -> Iterator[DocxParagraph | DocxTable]:
    """Yield Paragraph/Table objects in true document order (python-docx
    has no built-in mixed iterator).
    """
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield DocxParagraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield DocxTable(child, document)


def _paragraph_outline_level(paragraph: DocxParagraph) -> int | None:
    """Extract Word's built-in outline level from paragraph properties.
    One of several signals for clause-hierarchy detection; style name
    alone is unreliable across 3GPP document templates.
    """
    pPr = paragraph._p.pPr
    if pPr is None:
        return None
    outline_lvl = pPr.find(qn("w:outlineLvl"))
    if outline_lvl is None:
        return None
    val = outline_lvl.get(qn("w:val"))
    return int(val) if val is not None else None


def _paragraph_numbering_id(paragraph: DocxParagraph) -> int | None:
    pPr = paragraph._p.pPr
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    numId_el = numPr.find(qn("w:numId"))
    if numId_el is None:
        return None
    val = numId_el.get(qn("w:val"))
    return int(val) if val is not None else None


def _extract_table(table: DocxTable) -> RawTable:
    """Read a table's cells directly from OpenXML rather than
    `python-docx`'s `.cell()` grid, which silently repeats the same Cell
    object for merged spans and hides merge geometry.

    Word rows may omit grid columns at the start (`<w:gridBefore>`) or end
    (`<w:gridAfter>`), independent of `gridSpan`/`vMerge`. Real 3GPP
    tables (e.g. TS 24.501, TS 33.501, TS 29.244) have rows narrower than
    the header this way; each omitted column is represented as an empty
    placeholder cell so column indexing stays aligned with the true grid
    width.
    """
    rows: list[list[RawCell]] = []
    for row in table.rows:
        row_cells: list[RawCell] = []

        for _ in range(row.grid_cols_before):
            row_cells.append(RawCell(text="", grid_span=1, v_merge=None))

        for tc in row._tr.tc_lst:
            tcPr = tc.tcPr
            grid_span = 1
            v_merge: str | None = None
            if tcPr is not None:
                gridSpan_el = tcPr.find(qn("w:gridSpan"))
                if gridSpan_el is not None:
                    grid_span = int(gridSpan_el.get(qn("w:val")))
                vMerge_el = tcPr.find(qn("w:vMerge"))
                if vMerge_el is not None:
                    v_merge = vMerge_el.get(qn("w:val")) or "continue"
                    # OpenXML quirk: a bare <w:vMerge/> with no "val"
                    # attribute means "continue", not "restart".

            text = "\n".join(
                "".join(node.text or "" for node in p_el.iter(qn("w:t")))
                for p_el in tc.findall(qn("w:p"))
            ).strip()

            row_cells.append(RawCell(text=text, grid_span=grid_span, v_merge=v_merge))

        for _ in range(row.grid_cols_after):
            row_cells.append(RawCell(text="", grid_span=1, v_merge=None))

        rows.append(row_cells)
    return RawTable(rows=rows)


def parse_docx(path: Path) -> list[RawElement]:
    """Parse a DOCX file into an ordered list of raw paragraph/table
    elements, preserving document order.
    """
    logger.info("Parsing DOCX: {}", path.name)
    document = Document(str(path))

    elements: list[RawElement] = []
    paragraph_count = 0
    table_count = 0

    for item in _iter_block_items(document):
        if isinstance(item, DocxParagraph):
            elements.append(
                RawParagraph(
                    text=item.text,
                    style_name=item.style.name if item.style is not None else "",
                    outline_level=_paragraph_outline_level(item),
                    numbering_id=_paragraph_numbering_id(item),
                )
            )
            paragraph_count += 1
        elif isinstance(item, DocxTable):
            elements.append(_extract_table(item))
            table_count += 1

    logger.info(
        "Extracted {} structural elements ({} paragraphs, {} tables)",
        len(elements),
        paragraph_count,
        table_count,
    )
    return elements


def has_embedded_objects(path: Path) -> bool:
    """Detect embedded OLE objects or images via package relationships.
    Only presence is asserted, never contents.
    """
    document = Document(str(path))
    for rel in document.part.rels.values():
        if "image" in rel.reltype or "oleObject" in rel.reltype:
            return True
    return False
