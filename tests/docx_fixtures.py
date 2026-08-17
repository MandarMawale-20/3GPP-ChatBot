"""Builds a small, synthetic DOCX file that mimics the structural features
a real 3GPP specification exercises: numbered headings, a procedure with
steps, a table with both horizontal and vertical merges, an ASN.1 block in
a monospace style, and an annex.

Used across parser/structure/chunker/pipeline tests so none of them depend
on a real 3GPP file being present (SRD §24: synthetic fixtures, not the
real corpus, for unit tests).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE


def build_sample_docx(path: Path) -> None:
    document = Document()

    # Ensure a "TT" style exists — 3GPP's convention for monospaced/code
    # content, used here to mark the ASN.1 block (see ingestion/asn1_parser.py).
    styles = document.styles
    if "TT" not in [s.name for s in styles]:
        styles.add_style("TT", WD_STYLE_TYPE.PARAGRAPH)

    document.add_heading("5\tGeneral", level=1)
    document.add_paragraph("This clause describes general aspects of the procedure.")

    document.add_heading("5.1\tOverview", level=2)
    document.add_paragraph("The registration procedure allows the UE to register with the network.")

    document.add_heading("5.1.1\tRegistration procedure", level=3)
    document.add_paragraph("Step 1: UE sends REGISTRATION REQUEST to the AMF.")
    document.add_paragraph("Step 2: AMF processes the request and initiates authentication.")
    document.add_paragraph("NOTE: The exact authentication method depends on subscriber configuration.")

    document.add_heading("5.2\tParameters", level=2)
    table = document.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "IE"
    table.cell(0, 1).text = "Presence"
    table.cell(0, 2).text = "Description"
    table.cell(1, 0).text = "5GMM cause"
    table.cell(1, 1).text = "M"
    table.cell(1, 2).text = "Cause of the 5GMM rejection"
    table.cell(2, 1).text = "M"
    table.cell(2, 2).text = "Additional cause detail"
    # Vertical merge on column 0, rows 1-2 (simulates a repeated IE spanning
    # rows). Leave cell(2, 0) empty before merging — python-docx's merge()
    # concatenates both cells' paragraph content, so a non-empty value here
    # would duplicate text into the merged cell rather than producing the
    # clean single-value merge real 3GPP tables have.
    table.cell(1, 0).merge(table.cell(2, 0))

    document.add_heading("5.3\tASN.1 example", level=2)
    for line in [
        "RRCSetupRequest ::= SEQUENCE {",
        "    ue-Identity    InitialUE-Identity,",
        "    establishmentCause    EstablishmentCause,",
        "    spare    BIT STRING (SIZE (1))",
        "}",
    ]:
        p = document.add_paragraph(line)
        p.style = document.styles["TT"]

    document.add_heading("Annex A (normative): Change history", level=1)
    document.add_paragraph("This annex records the change history of the present document.")

    document.save(str(path))
